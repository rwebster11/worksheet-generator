import os
import io
import re # For parsing YouTube URL
import random # Needed for word search logic
from copy import deepcopy # Needed for word search logic
import traceback # Import traceback for detailed error logging
import logging # Import standard logging
from datetime import datetime, timezone # Added timezone for UTC awareness
from PIL import Image, ImageDraw, ImageFont
import math
# Third-party imports
import anthropic
import httpx # Still needed by Anthropic client internally, even without proxy
from dotenv import load_dotenv
from flask import Flask, request, jsonify, send_from_directory, send_file
from flask_sqlalchemy import SQLAlchemy
from bs4 import BeautifulSoup
import docx
from docx.shared import Inches, Pt
from youtube_transcript_api import YouTubeTranscriptApi, TranscriptsDisabled, NoTranscriptFound # Keep for now
import json
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import uuid # For generating unique filenames
from werkzeug.utils import secure_filename # For basic filename security (optional but good practice)
from PIL import Image, ImageDraw, ImageFont # For DOCX image generation
import io # For handling image data in memory
import json
# Configure basic logging early
logging.basicConfig(level=logging.INFO, format='%(asctime)s %(levelname)s:%(name)s:%(message)s')

# Load environment variables from .env file
logging.info("Attempting to load .env file...")
dotenv_path = os.path.join(os.path.dirname(__file__), '.env') # Explicit path
found_dotenv = load_dotenv(dotenv_path=dotenv_path, verbose=True) # Be verbose
logging.info(f".env file found and loaded: {found_dotenv}")
# Removed duplicate load_dotenv()

# Initialize the Flask application
app = Flask(__name__)

# --- Database Configuration ---
basedir = os.path.abspath(os.path.dirname(__file__))
db_path = os.path.join(basedir, 'worksheet_library.db')
app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{db_path}'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
# --- End Database Configuration ---
# --- Upload Folder Configuration --- ADD THIS SECTION ---
UPLOAD_FOLDER = os.path.join(basedir, 'static', 'uploads')
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
    logging.info(f"Created upload folder at: {UPLOAD_FOLDER}")
else:
    logging.info(f"Upload folder exists at: {UPLOAD_FOLDER}")
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024 # Example: 16MB upload limit
# --- End Upload Folder Configuration ---
# Initialize the SQLAlchemy extension
db = SQLAlchemy(app)
with app.app_context():
    db.create_all()
    logging.info("Database tables checked/created (if they didn't exist).")
# --- Database Model Definitions ---

# Association Table: Links Worksheets and GeneratedItems, storing order
worksheet_items_association = db.Table('worksheet_items_association',
    db.metadata,
    db.Column('worksheet_id', db.Integer, db.ForeignKey('worksheet.id'), primary_key=True),
    db.Column('generated_item_id', db.Integer, db.ForeignKey('generated_item.id'), primary_key=True), # Matches your column name
    db.Column('item_order', db.Integer, nullable=False)
)

# GeneratedItem Model
class GeneratedItem(db.Model):
    __tablename__ = 'generated_item'
    id = db.Column(db.Integer, primary_key=True)
    item_type = db.Column(db.String(50), nullable=False)
    source_topic = db.Column(db.String(250), nullable=True)
    source_url = db.Column(db.String(500), nullable=True)
    grade_level = db.Column(db.String(50), nullable=False)
    content_html = db.Column(db.Text, nullable=False)
    item_data_json = db.Column(db.Text, nullable=True)
    creation_date = db.Column(db.DateTime(timezone=True), nullable=False, default=lambda: datetime.now(timezone.utc))
    last_modified_date = db.Column(db.DateTime(timezone=True), nullable=False, default=lambda: datetime.now(timezone.utc), onupdate=lambda: datetime.now(timezone.utc))

    worksheets = db.relationship(
        'Worksheet',
        secondary='worksheet_items_association',
        back_populates='items',
        lazy='selectin'
    )

    def __repr__(self):
        return f'<GeneratedItem id={self.id} type={self.item_type}>'

# Worksheet Model
class Worksheet(db.Model):
    __tablename__ = 'worksheet'
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(250), nullable=False, default="Untitled Worksheet")
    creation_date = db.Column(db.DateTime(timezone=True), nullable=False, default=lambda: datetime.now(timezone.utc))
    last_modified_date = db.Column(db.DateTime(timezone=True), nullable=False, default=lambda: datetime.now(timezone.utc), onupdate=lambda: datetime.now(timezone.utc))

    items = db.relationship(
        'GeneratedItem',
        secondary='worksheet_items_association',
        order_by=worksheet_items_association.c.item_order,
        back_populates='worksheets',
        lazy='selectin'
    )

    def __repr__(self):
        return f'<Worksheet id={self.id} title="{self.title}">'

# Configure mappers after models are defined
db.configure_mappers()

# --- Anthropic Client Initialization ---
client = None
try:
    logging.info("Attempting to initialize Anthropic client...")
    api_key_from_env = os.getenv("ANTHROPIC_API_KEY")
    if not api_key_from_env:
        logging.error("CRITICAL: ANTHROPIC_API_KEY not found in environment variables.")
    else:
        logging.info("ANTHROPIC_API_KEY found. Initializing client...")
        # Initialize without unused proxy setup
        client = anthropic.Anthropic(api_key=api_key_from_env)
        logging.info("Anthropic client object CREATED successfully.")
except Exception as e:
    logging.error(f"CRITICAL: Exception during Anthropic client initialization: {e}", exc_info=True)
    client = None

if client is None:
    logging.warning("Anthropic client is None after initialization attempt.")
else:
    logging.info("Anthropic client appears to be initialized.")

ANTHROPIC_MODEL_NAME = "claude-3-opus-20240229"

# --- Prompt Creation Functions ---
def create_mcq_prompt(topic, grade_level, num_questions=5):
    # ... (keep function as is) ...
    prompt = f"""You are an expert curriculum developer creating assessment questions.
Generate {num_questions} multiple-choice questions (MCQs) about the topic '{topic}'.
The target audience is {grade_level} students.

**CRITICAL INSTRUCTIONS for Output Format:**
1.  **Question Generation:** Create {num_questions} distinct MCQs covering different important aspects of the topic. Adapt vocabulary and complexity for the {grade_level}.
2.  **Options:** For each question, provide exactly four options: one correct answer (labeled A, B, C, or D) and three plausible distractors.
3.  **Formatting:**
    *   Present each question clearly, starting with a number (e.g., "1.", "2.").
    *   List the options below each question, labeled "A.", "B.", "C.", "D.".
    *   **DO NOT** indicate the correct answer within the question/option list itself.
4.  **Answer Key:** After ALL questions and their options are listed, provide a separate section titled exactly "Answer Key:".
5.  **Key Format:** In the Answer Key section, list the question number and the correct option letter (e.g., "1. C", "2. A", "3. B", etc.), each on a new line.
6.  **Strict Output:** Output ONLY the numbered questions with their A/B/C/D options, followed by the "Answer Key:" section and the key itself. No extra text, introductions, or explanations.

Topic: {topic}
Audience: {grade_level}

MCQs:
"""
    return prompt

def create_word_list_prompt(topic, grade_level, num_words=15):
    # ... (keep function as is) ...
    prompt = f"""You are a helpful vocabulary assistant.
Generate a list of exactly {num_words} single keywords or very short (2-word max) key phrases relevant to the topic '{topic}' for a {grade_level} audience.

**Instructions:**
1.  Focus on specific nouns, verbs, or essential terms related to '{topic}'.
2.  Ensure words are appropriate for the {grade_level}.
3.  Provide *only* the list of words/phrases.
4.  Format the output as a numbered list (e.g., "1. Photosynthesis", "2. Chlorophyll", "3. Carbon Dioxide").
5.  Do not include definitions, explanations, or any text other than the numbered list.

Topic: {topic}
Audience: {grade_level}
Number of Words: {num_words}

Keyword List:
"""
    return prompt

def create_text_block_prompt(topic, grade_level, focus=None):
    # ... (keep function as is) ...
    focus_instruction = f"Focus on: {focus}" if focus else "Provide a general overview."
    prompt = f"""You are a helpful assistant generating informational content.
Write a clear and concise text block about the topic '{topic}'.
The target audience is {grade_level} students.
{focus_instruction}
Aim for 1-3 paragraphs unless specified otherwise.
Ensure the language is appropriate for the specified grade level.

**Output ONLY the text block itself.** No titles, introductions, or extra formatting other than paragraphs.

Topic: {topic}
Audience: {grade_level}

Text Block:
"""
    return prompt

def create_true_false_prompt(topic, grade_level, num_statements=8):
    # ... (keep function as is) ...
    prompt = f"""You are an expert educator creating assessment materials.
Generate {num_statements} True/False statements about the core concepts of the topic '{topic}'.
The target audience is {grade_level} students.

**CRITICAL INSTRUCTIONS:**
1.  **Statement Generation:** Create {num_statements} clear statements. Ensure a mix of reasonably challenging statements that are definitively TRUE and definitively FALSE based on common knowledge for the {grade_level} audience regarding the topic '{topic}'. Avoid ambiguity or opinion-based statements.
2.  **Formatting:** Present each statement clearly, starting with a number (e.g., "1.", "2.").
3.  **Answer Key:** After ALL statements are listed, provide a separate section titled exactly "Answer Key:".
4.  **Key Format:** In the Answer Key section, list the statement number followed by the word "True" or "False" (e.g., "1. True", "2. False", "3. True", etc.), each on a new line.
5.  **Strict Output:** Output ONLY the numbered statements, followed by the "Answer Key:" section and the key itself. No extra text, introductions, or explanations.

Topic: {topic}
Audience: {grade_level}

True/False Statements:
"""
    return prompt

def create_gap_fill_prompt(topic, grade_level="middle school", num_sentences=7):
    # ... (keep function as is) ...
    prompt = f"""You are an expert teacher creating educational resources.
Generate a gap-fill (cloze) activity about the topic '{topic}'.
The target audience is {grade_level} students.

**Instructions:**
1.  Create exactly {num_sentences} unique sentences. Each sentence **MUST** focus on a **different** important aspect or key term of the topic '{topic}'. Do not just rephrase the same core idea.
2.  In each sentence, identify the single most important key term or concept specific to that sentence's point and replace it with '_________'.
3.  **CRITICAL REQUIREMENT: Each blank MUST be fillable with a DIFFERENT word.** Do NOT reuse the same answer word for multiple blanks. The goal is to test a variety of key vocabulary related to the topic. Using the exact same word in two or more blanks is forbidden.
4.  After the sentences, provide a numbered list titled 'Answer Key:' that clearly lists the single word removed for each blank in the correct order.
5.  Output ONLY the worksheet sentences and the Answer Key. No introductory phrases, explanations, conversational text, or titles other than 'Answer Key:'.

Topic: {topic}
Audience: {grade_level}
Number of Sentences: {num_sentences}

Worksheet:
"""
    return prompt

def extract_video_id(url): # Keep for now
    # ... (keep function as is) ...
    patterns = [
        r'(?:https?:\/\/)?(?:www\.)?youtube\.com\/watch\?v=([a-zA-Z0-9_-]{11})',
        r'(?:https?:\/\/)?(?:www\.)?youtu\.be\/([a-zA-Z0-9_-]{11})',
        r'(?:https?:\/\/)?(?:www\.)?youtube\.com\/embed\/([a-zA-Z0-9_-]{11})',
        r'(?:https?:\/\/)?(?:www\.)?youtube\.com\/v\/([a-zA-Z0-9_-]{11})',
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match: return match.group(1)
    return None

def create_comprehension_prompt(transcript_text, num_questions=5): # Keep for now
    # ... (keep function as is) ...
    prompt = f"""You are an expert educator. Based on the following video transcript, please generate {num_questions} insightful comprehension questions that test understanding of the key information presented.

**Instructions for Questions:**
1.  Ensure questions cover different aspects of the transcript (e.g., main ideas, specific details, inferences if possible).
2.  Phrase questions clearly and concisely.
3.  The questions should be suitable for someone who has just watched the video (assume context from the transcript).
4.  Do NOT provide the answers to the questions.
5.  Format the output as a numbered list.
6.  Output ONLY the numbered list of questions, nothing else (no preamble like "Here are the questions:").

**Video Transcript:**
--- START TRANSCRIPT ---
{transcript_text}
--- END TRANSCRIPT ---

**{num_questions} Comprehension Questions:**
"""
    return prompt

def create_pasted_comprehension_prompt(pasted_text, num_questions=5):
    # ... (keep function as is) ...
    max_chars = 15000
    truncated_text = pasted_text[:max_chars]
    if len(pasted_text) > max_chars: print(f"Warning: Pasted text truncated to {max_chars} chars for AI prompt.")
    prompt = f"""You are an expert educator designing reading comprehension assessments.
Based *only* on the following provided text passage, generate {num_questions} insightful comprehension questions that test understanding of the key information presented within that text.
Also provide a brief answer key outlining the expected points or a model answer for each question based *only* on the provided text.

**CRITICAL INSTRUCTIONS:**
1.  **Question Generation:** Create {num_questions} distinct questions directly related to the content of the provided text passage. Questions should require recall, explanation, or analysis *of the passage*.
2.  **Question Formatting:** Present each question clearly, starting with a number (e.g., "1.", "2.").
3.  **Answer Key Generation:** After ALL questions are listed, provide a separate section titled exactly "Answer Key:".
4.  **Key Content:** For each question number in the Answer Key section, provide key points or a model answer derived *solely from the provided passage*.
5.  **Key Formatting:** Start each answer key item with the corresponding question number (e.g., "1.", "2.").
6.  **Strict Output:** Output ONLY the numbered list of questions, followed by the "Answer Key:" section and the key itself. No extra text, introductions, or commentary about the text quality.

**Provided Text Passage:**
--- START TEXT ---
{truncated_text}
--- END TEXT ---

**{num_questions} Comprehension Questions:**
[Generate Questions Here Following Format]

Answer Key:
[Generate Answer Key Here Following Format]
"""
    return prompt

def create_short_answer_prompt(topic, grade_level, num_questions=5):
    # ... (keep function as is) ...
    prompt = f"""You are an expert educator designing formative assessments.
Generate {num_questions} open-ended short answer questions about the key aspects of the topic '{topic}'.
Also provide a brief answer key outlining the expected points or a model answer for each question.
The target audience is {grade_level} students.

**CRITICAL INSTRUCTIONS:**
1.  **Question Generation:** Create {num_questions} distinct questions that require students to recall, explain, or briefly analyze information related to '{topic}'. Questions should encourage answers of 1-3 sentences. Adapt vocabulary and complexity for the {grade_level}.
2.  **Question Formatting:** Present each question clearly, starting with a number (e.g., "1.", "2."). List each question on a new line.
3.  **Answer Key Generation:** After ALL questions are listed, provide a separate section titled exactly "Answer Key:".
4.  **Key Content:** For each question number in the Answer Key section, provide either:
    *   A bulleted list of the key facts/points expected in a good answer.
    *   OR a brief model answer (1-2 sentences).
    Keep the answer key concise and focused on grading criteria.
5.  **Key Formatting:** Start each answer key item with the corresponding question number (e.g., "1.", "2.").
6.  **Strict Output:** Output ONLY the numbered list of questions, followed by the "Answer Key:" section and the key itself. No extra text, introductions, or explanations.

Topic: {topic}
Audience: {grade_level}

Short Answer Questions:
[Generate Questions Here Following Format]

Answer Key:
[Generate Answer Key Here Following Format]
"""
    return prompt

def create_similar_questions_prompt(example_question, num_questions=3, grade_level="Not Specified"):
    # ... (keep function as is) ...
    prompt = f"""
You are an expert question designer replicating educational assessment items. Your task is to analyze the provided example question, identify the core concept and calculation type being tested, and then generate {num_questions} NEW questions that assess the SAME core concept at a SIMILAR difficulty level suitable for {grade_level} students.

**Analysis of Example:**
First, briefly state the core concept or skill tested by the example question. (e.g., "Calculating speed given distance and time", "Solving a two-step linear equation", "Calculating molarity").

**Instructions for New Questions:**
1.  **Maintain Core Concept:** Each new question MUST test the exact same fundamental principle or calculation type as the example.
2.  **Maintain Difficulty:** The complexity of the numbers used, the number of steps required, and the required conceptual understanding should closely match the example question and be appropriate for the target {grade_level}.
3.  **Construct Realistic Scenarios:** For each new question:
    a.  **Determine a realistic target answer** appropriate for the concept and grade level (e.g., a percentage yield between 60-95%, a reasonable speed for a car).
    b.  **Work Backwards:** Choose plausible input numbers (like initial mass, theoretical yield, distance, time, etc.) that will mathematically lead to the target answer you determined. Ensure these input numbers are also contextually sensible (e.g., realistic masses, times, distances).
    c.  **Vary Context & Specifics:** Create a genuinely new scenario (different chemicals, objects, situations) using these generated numbers. You can vary units where appropriate (e.g., grams to kg) and potentially the specific variable being solved for (if applicable), while keeping the underlying problem type identical to the example. Do NOT just rephrase the example question.
4.  **Formatting:** Present each new question clearly, starting with a number (e.g., "1.", "2."). List each question on a new line.
5.  **Answer Key Generation:** After ALL new questions are listed, provide a separate section titled exactly "Answer Key:".
6. Key Content: In the Answer Key section, for each question number, show the key calculation steps required to solve the problem, followed by the final numerical answer (including units if appropriate). Make the steps clear and concise.

    Example Format:
    1. Steps:
       Step 1: Formula (e.g., Speed = Distance / Time)
       Step 2: Substitute values (e.g., Speed = 300 miles / 5 hours)
       Step 3: Calculate (e.g., Speed = 60)
       Final Answer: 60 mph
    2. Steps:
       ...
       Final Answer: ...

7.  **Strict Output:** Output ONLY the brief "Analysis of Example", the numbered list of new questions, the "Answer Key:" title, and the key itself. No extra text, introductions, or explanations.

**Example Question Provided:**
--- START EXAMPLE ---
{example_question}
--- END EXAMPLE ---

**Analysis of Example:**
[AI writes analysis here]

**{num_questions} New Similar Questions:**
[AI generates new questions here]

Answer Key:
[AI generates answer key for the NEW questions here]
"""
    return prompt

# --- Word Search Generation Logic (Adapted from make_wordsearch.py) ---

# Constants
NMAX_GRID = 32
ALPHABET = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'

# --- Masking Functions ---
def circle_mask(grid, nrows, ncols):
    r2 = min(ncols, nrows)**2 // 4
    cx, cy = ncols//2, nrows // 2
    for irow in range(nrows):
        for icol in range(ncols):
            if (irow - cy)**2 + (icol - cx)**2 > r2:
                grid[irow][icol] = '*'
def squares_mask(grid, nrows, ncols):
    a = int(0.38 * min(ncols, nrows))
    cy = nrows // 2; cx = ncols // 2
    for irow in range(nrows):
        for icol in range(ncols):
            if a <= icol < ncols-a:
                if irow < cy-a or irow > cy+a: grid[irow][icol] = '*'
            if a <= irow < nrows-a:
                if icol < cx-a or icol > cx+a: grid[irow][icol] = '*'
def no_mask(grid, nrows, ncols): pass
apply_mask = {None: no_mask, 'circle': circle_mask, 'squares': squares_mask}

# --- Grid Creation and Filling ---
def make_initial_grid(nrows, ncols, mask_name=None):
    if mask_name not in apply_mask:
        logging.warning(f"Unknown mask name '{mask_name}'. Using no mask.")
        mask_name = None
    grid = [[' ']*ncols for _ in range(nrows)]
    apply_mask[mask_name](grid, nrows, ncols)
    return grid

def fill_grid_randomly(grid, nrows, ncols):
    for irow in range(nrows):
        for icol in range(ncols):
            if grid[irow][icol] == ' ': grid[irow][icol] = random.choice(ALPHABET)

def remove_mask_chars(grid, nrows, ncols):
    for irow in range(nrows):
        for icol in range(ncols):
            if grid[irow][icol] == '*': grid[irow][icol] = ' '

# --- Word Placement Logic ---
def _try_place_words(nrows, ncols, wordlist, allow_backwards_words=True, mask_name=None):
    grid = make_initial_grid(nrows, ncols, mask_name)

    def test_candidate(grid, irow, icol, dx, dy, word):
        word_len = len(word)
        for j in range(word_len):
            test_row, test_col = irow + j*dy, icol + j*dx
            if not (0 <= test_row < nrows and 0 <= test_col < ncols): return False
            if grid[test_row][test_col] not in (' ', word[j]): return False
        return True

    def place_word_in_grid(grid, word):
        word = ''.join(filter(str.isalnum, word)).upper()
        if not word: return None
        dxdy_choices = [(0,1), (1,0), (1,1), (1,-1)]; random.shuffle(dxdy_choices)

        for (dx, dy) in dxdy_choices:
            word_to_place = word
            if allow_backwards_words and random.choice([True, False]): word_to_place = word[::-1]

            n = len(word_to_place)
            colmin = 0; colmax = ncols - n if dx != 0 else ncols - 1
            rowmin = 0 if dy >= 0 else n - 1; rowmax = nrows - n if dy > 0 else nrows - 1
            if dy < 0: rowmax = nrows -1 # Correct range for up-diagonals

            if colmax < colmin or rowmax < rowmin: continue

            candidates = []
            for r in range(rowmin, rowmax + 1):
                 for c in range(colmin, colmax + 1):
                      end_r, end_c = r + (n-1)*dy, c + (n-1)*dx
                      if 0 <= end_r < nrows and 0 <= end_c < ncols:
                           if test_candidate(grid, r, c, dx, dy, word_to_place): candidates.append((r, c))

            if not candidates: continue
            start_row, start_col = random.choice(candidates)
            current_row, current_col = start_row, start_col
            for char in word_to_place:
                grid[current_row][current_col] = char
                current_row += dy; current_col += dx
            logging.debug(f"Placed '{word}' at ({start_row},{start_col}) orientation ({dx},{dy})")
            return True # Indicate success

        logging.warning(f"Failed to place word: {word}")
        return False

    successfully_placed_words = []
    sorted_wordlist = sorted([w for w in wordlist if w], key=len, reverse=True)
    for word in sorted_wordlist:
        if place_word_in_grid(grid, word): successfully_placed_words.append(word.upper())
        else: return None, None

    solution_grid = deepcopy(grid); fill_grid_randomly(grid, nrows, ncols)
    remove_mask_chars(grid, nrows, ncols); remove_mask_chars(solution_grid, nrows, ncols)
    return grid, sorted(successfully_placed_words)

# --- Top-Level Generation Function ---
class PuzzleGenerationError(Exception): pass

def generate_wordsearch_grid(wordlist, nrows, ncols, allow_backwards_words=True, mask_name=None, attempts=10):
    if nrows > NMAX_GRID or ncols > NMAX_GRID: raise ValueError(f'Max grid dimension is {NMAX_GRID}')
    if not wordlist: raise ValueError("Word list cannot be empty.")

    max_dimension = max(nrows, ncols); cleaned_wordlist = []
    for word in wordlist:
         cleaned_word = ''.join(filter(str.isalnum, word)).upper()
         if cleaned_word:
              if len(cleaned_word) > max_dimension: raise ValueError(f"Word '{cleaned_word}' > {max_dimension} chars (grid {nrows}x{ncols})")
              cleaned_wordlist.append(cleaned_word)
    if not cleaned_wordlist: raise ValueError("No valid words after cleaning.")

    for i in range(attempts):
        logging.info(f"Word search generation attempt {i+1}/{attempts}...")
        grid, placed_words = _try_place_words(nrows, ncols, cleaned_wordlist, allow_backwards_words, mask_name)
        if grid is not None:
            logging.info(f"Successfully generated grid in {i+1} attempt(s). Placed {len(placed_words)} words.")
            return grid, placed_words
    raise PuzzleGenerationError(f"Failed to place all words after {attempts} attempts.")
# --- End of Word Search Generation Logic ---
def create_wordsearch_image(grid_list, cell_size=30, font_size=18, font_path=None, grid_line_padding=1):
    """Creates a PIL Image object of the word search grid with complete borders."""
    if not grid_list or not grid_list[0]:
        logging.warning("create_wordsearch_image: Empty grid_list provided.")
        return None

    nrows = len(grid_list)
    ncols = len(grid_list[0])
    
    # Calculate actual grid dimensions
    grid_actual_width = ncols * cell_size
    grid_actual_height = nrows * cell_size

    # Image canvas size includes padding for borders
    img_width = grid_actual_width + (2 * grid_line_padding)
    img_height = grid_actual_height + (2 * grid_line_padding)

    image = Image.new('RGB', (img_width, img_height), 'white')
    draw = ImageDraw.Draw(image)

    try:
        font_to_use = ImageFont.truetype(font_path or "arial.ttf", font_size)
    except IOError:
        logging.warning(f"Word search font not found at '{font_path or "arial.ttf"}'. Using default.")
        font_to_use = ImageFont.load_default()

    # Draw letters and cell backgrounds/internal lines
    for r in range(nrows):
        for c in range(ncols):
            char_to_draw = grid_list[r][c]
            
            # Top-left corner of the cell on the canvas (including padding)
            cell_x0_canvas = c * cell_size + grid_line_padding
            cell_y0_canvas = r * cell_size + grid_line_padding
            
            # Bottom-right corner of the cell on the canvas
            # cell_x1_canvas = cell_x0_canvas + cell_size
            # cell_y1_canvas = cell_y0_canvas + cell_size

            # Draw letter centered in cell
            if char_to_draw and char_to_draw.strip(): # Only draw if there's a character
                try:
                    # Use textbbox for better centering if available (Pillow 9.2.0+)
                    bbox = draw.textbbox((0,0), char_to_draw, font=font_to_use, anchor="lt") # Left-Top anchor for bbox calc
                    text_width = bbox[2] - bbox[0]
                    text_height = bbox[3] - bbox[1] # This includes ascender/descender
                    # For more precise vertical centering, consider font metrics
                    ascent, descent = font_to_use.getmetrics()
                    actual_text_visual_height = ascent # Height above baseline for typical cap height

                    text_draw_x = cell_x0_canvas + (cell_size - text_width) / 2
                    text_draw_y = cell_y0_canvas + (cell_size - actual_text_visual_height) / 2 # Center based on ascent
                    
                    draw.text((text_draw_x, text_draw_y), char_to_draw, fill='black', font=font_to_use)
                except (AttributeError, TypeError): # Fallback for older Pillow or if anchor not supported well
                    text_width, text_height = draw.textsize(char_to_draw, font=font_to_use) # Deprecated
                    text_draw_x = cell_x0_canvas + (cell_size - text_width) / 2
                    text_draw_y = cell_y0_canvas + (cell_size - text_height) / 2
                    draw.text((text_draw_x, text_draw_y), char_to_draw, fill='black', font=font_to_use)

    # Draw grid lines (horizontal and vertical) AFTER letters to ensure they are on top of any cell bg
    # Horizontal lines
    for r in range(nrows + 1):
        y = r * cell_size + grid_line_padding
        draw.line([(grid_line_padding, y), (grid_actual_width + grid_line_padding, y)], fill='grey', width=1)

    # Vertical lines
    for c in range(ncols + 1):
        x = c * cell_size + grid_line_padding
        draw.line([(x, grid_line_padding), (x, grid_actual_height + grid_line_padding)], fill='grey', width=1)

    return image
# --- Standard API Routes ---
# (Generic Error Handlers - Can be refactored later)
def handle_anthropic_error(e):
    if isinstance(e, anthropic.APIConnectionError): code=503; msg=f'AI Connection Error: {e}'
    elif isinstance(e, anthropic.RateLimitError): code=429; msg='AI Rate Limit Exceeded.'
    elif isinstance(e, anthropic.AuthenticationError): code=401; msg=f'AI Authentication Error: {e}'
    elif isinstance(e, anthropic.APIStatusError):
        code = e.status_code; msg = f'AI service error (Status {code})'
        try: error_details = e.response.json(); msg += f": {error_details.get('error', {}).get('message', e.response.text)}"
        except Exception: msg += f": {e.response.text}"
    else: code=500; msg='An unexpected error occurred during AI call.'
    logging.error(f"{msg} - Exception: {e}", exc_info=(code == 500)) # Log traceback for unexpected
    return jsonify({'status': 'error', 'message': msg}), code

# --- List Saved Items Route ---
@app.route('/list_items', methods=['GET'])
def list_items_route():
    logging.info("Received request at /list_items")
    try:
        items = GeneratedItem.query.order_by(GeneratedItem.last_modified_date.desc()).limit(20).all()
        items_list = [{'id': item.id, 'item_type': item.item_type, 'source_topic': item.source_topic,
                       'source_url': item.source_url, 'grade_level': item.grade_level,
                       'last_modified': item.last_modified_date.isoformat()} for item in items]
        return jsonify({'status': 'success', 'items': items_list})
    except Exception as e:
        logging.error(f"Error retrieving items from database: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': 'Error retrieving items from library.'}), 500

@app.route('/upload_base_image', methods=['POST'])
def upload_base_image():
    logging.info("Received request at /upload_base_image")
    if 'imageFile' not in request.files:
        logging.warning("No 'imageFile' part in the request files.")
        return jsonify({'status': 'error', 'message': 'No image file part in the request'}), 400

    file = request.files['imageFile']
    if file.filename == '':
        logging.warning("'imageFile' present but no selected file.")
        return jsonify({'status': 'error', 'message': 'No selected image file'}), 400

    # Basic file type validation (check MIME type or extension)
    allowed_extensions = {'png', 'jpg', 'jpeg', 'gif', 'webp'}
    file_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
    if not file_ext or file_ext not in allowed_extensions:
        logging.warning(f"Disallowed file extension: {file_ext}")
        return jsonify({'status': 'error', 'message': f'Invalid image format. Allowed: {", ".join(allowed_extensions)}'}), 400

    try:
        # Generate a unique filename using UUID
        # Keep original extension for browser compatibility/MIME type detection
        unique_filename = str(uuid.uuid4()) + '.' + file_ext
        # Using secure_filename is still good practice to avoid path traversal issues etc.
        # but we discard its output as we use UUID for uniqueness.
        _ = secure_filename(file.filename) # Run it for validation side-effects
        save_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        logging.info(f"Attempting to save uploaded image to: {save_path}")

        file.save(save_path)
        logging.info(f"Image saved successfully as {unique_filename}")

        image_url = f"/static/uploads/{unique_filename}" # URL path for the frontend

        return jsonify({
            'status': 'success',
            'filename': unique_filename,
            'image_url': image_url
        })

    except Exception as e:
        logging.error(f"Error saving uploaded image: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': 'Server error during image upload.'}), 500

# --- Generate Routes ---
@app.route('/generate_true_false', methods=['POST'])
def generate_true_false_route():
    logging.info("Received request at /generate_true_false")
    if client is None: return jsonify({'status': 'error', 'message': 'Server error: AI client not initialized'}), 500
    if not request.is_json: return jsonify({'status': 'error', 'message': 'Request must be JSON'}), 400
    try:
        data = request.get_json(); topic = data.get('topic'); grade_level = data.get('grade_level', 'middle school'); num_statements_req = data.get('num_statements', 8)
        if not topic: return jsonify({'status': 'error', 'message': 'Missing "topic"'}), 400
        try: num_statements = int(num_statements_req); assert 3 <= num_statements <= 20
        except (ValueError, TypeError, AssertionError): return jsonify({'status': 'error', 'message': 'Invalid num_statements (must be int 3-20)'}), 400
        logging.info(f"T/F request: Topic='{topic}', Grade='{grade_level}', NumStatements={num_statements}")
        prompt_content = create_true_false_prompt(topic, grade_level, num_statements)
        message = client.messages.create(model=ANTHROPIC_MODEL_NAME, max_tokens=1000 + (num_statements * 50), temperature=0.6, messages=[{ "role": "user", "content": prompt_content }])
        if message.content and len(message.content) > 0 and hasattr(message.content[0], 'text'): generated_content = message.content[0].text
        else: raise ValueError("Failed to parse content from AI response.")
        logging.info(f"Generated T/F content length: {len(generated_content)} chars")
        return jsonify({'status': 'success', 'true_false_content': generated_content.strip()})
    except ValueError as ve: logging.warning(f"Validation Error: {ve}"); return jsonify({'status': 'error', 'message': str(ve)}), 400
    except (anthropic.APIError, httpx.RequestError) as e: return handle_anthropic_error(e)
    except Exception as e: logging.error(f"Unexpected error in /generate_true_false: {e}", exc_info=True); return jsonify({'status': 'error', 'message': 'Internal server error.'}), 500

@app.route('/generate_mcq', methods=['POST'])
def generate_mcq_route():
    logging.info("Received request at /generate_mcq")
    if client is None: return jsonify({'status': 'error', 'message': 'Server error: AI client not initialized'}), 500
    if not request.is_json: return jsonify({'status': 'error', 'message': 'Request must be JSON'}), 400
    try:
        data = request.get_json(); topic = data.get('topic'); grade_level = data.get('grade_level', 'middle school'); num_questions_req = data.get('num_questions', 5)
        if not topic: return jsonify({'status': 'error', 'message': 'Missing "topic"'}), 400
        try: num_questions = int(num_questions_req); assert 2 <= num_questions <= 15
        except (ValueError, TypeError, AssertionError): return jsonify({'status': 'error', 'message': 'Invalid num_questions (must be int 2-15)'}), 400
        logging.info(f"MCQ request: Topic='{topic}', Grade='{grade_level}', NumQuestions={num_questions}")
        prompt_content = create_mcq_prompt(topic, grade_level, num_questions)
        message = client.messages.create(model=ANTHROPIC_MODEL_NAME, max_tokens=1500 + (num_questions * 100), temperature=0.7, messages=[{ "role": "user", "content": prompt_content }])
        if message.content and len(message.content) > 0 and hasattr(message.content[0], 'text'): generated_content = message.content[0].text
        else: raise ValueError("Failed to parse content from AI response.")
        logging.info(f"Generated MCQ content length: {len(generated_content)} chars")
        return jsonify({'status': 'success', 'mcq_content': generated_content.strip()})
    except ValueError as ve: logging.warning(f"Validation Error: {ve}"); return jsonify({'status': 'error', 'message': str(ve)}), 400
    except (anthropic.APIError, httpx.RequestError) as e: return handle_anthropic_error(e)
    except Exception as e: logging.error(f"Unexpected error in /generate_mcq: {e}", exc_info=True); return jsonify({'status': 'error', 'message': 'Internal server error.'}), 500

@app.route('/generate_text_block', methods=['POST'])
def generate_text_block_route():
    logging.info("Received request at /generate_text_block")
    if client is None: return jsonify({'status': 'error', 'message': 'Server error: AI client not initialized'}), 500
    if not request.is_json: return jsonify({'status': 'error', 'message': 'Request must be JSON'}), 400
    try:
        data = request.get_json(); topic = data.get('topic'); grade_level = data.get('grade_level', 'middle school'); focus = data.get('focus')
        if not topic: return jsonify({'status': 'error', 'message': 'Missing "topic"'}), 400
        logging.info(f"Text Block request: Topic='{topic}', Grade='{grade_level}', Focus='{focus}'")
        prompt_content = create_text_block_prompt(topic, grade_level, focus)
        message = client.messages.create(model=ANTHROPIC_MODEL_NAME, max_tokens=800, temperature=0.7, messages=[{ "role": "user", "content": prompt_content }])
        if message.content and len(message.content) > 0 and hasattr(message.content[0], 'text'): generated_content = message.content[0].text
        else: raise ValueError("Failed to parse content from AI response.")
        logging.info(f"TextBlock AI: Generated content for text block: '{generated_content[:200]}...'")
        return jsonify({'status': 'success', 'text_block_content': generated_content.strip()})
    except ValueError as ve: logging.warning(f"Validation Error: {ve}"); return jsonify({'status': 'error', 'message': str(ve)}), 400
    except (anthropic.APIError, httpx.RequestError) as e: return handle_anthropic_error(e)
    except Exception as e: logging.error(f"Unexpected error in /generate_text_block: {e}", exc_info=True); return jsonify({'status': 'error', 'message': 'Internal server error.'}), 500

@app.route('/generate_short_answer', methods=['POST'])
def generate_short_answer_route():
    logging.info("Received request at /generate_short_answer")
    if client is None: return jsonify({'status': 'error', 'message': 'Server error: AI client not initialized'}), 500
    if not request.is_json: return jsonify({'status': 'error', 'message': 'Request must be JSON'}), 400
    try:
        data = request.get_json(); topic = data.get('topic'); grade_level = data.get('grade_level', 'middle school'); num_questions_req = data.get('num_questions', 5)
        if not topic: return jsonify({'status': 'error', 'message': 'Missing "topic"'}), 400
        try: num_questions = int(num_questions_req); assert 2 <= num_questions <= 10
        except (ValueError, TypeError, AssertionError): return jsonify({'status': 'error', 'message': 'Invalid num_questions (must be int 2-10)'}), 400
        logging.info(f"SA request: Topic='{topic}', Grade='{grade_level}', NumQuestions={num_questions}")
        prompt_content = create_short_answer_prompt(topic, grade_level, num_questions)
        message = client.messages.create(model=ANTHROPIC_MODEL_NAME, max_tokens=1200 + (num_questions * 100), temperature=0.7, messages=[{ "role": "user", "content": prompt_content }])
        if message.content and len(message.content) > 0 and hasattr(message.content[0], 'text'): generated_content = message.content[0].text
        else: raise ValueError("Failed to parse content from AI response.")
        logging.info(f"Generated SA content length: {len(generated_content)} chars")
        return jsonify({'status': 'success', 'short_answer_content': generated_content.strip()})
    except ValueError as ve: logging.warning(f"Validation Error: {ve}"); return jsonify({'status': 'error', 'message': str(ve)}), 400
    except (anthropic.APIError, httpx.RequestError) as e: return handle_anthropic_error(e)
    except Exception as e: logging.error(f"Unexpected error in /generate_short_answer: {e}", exc_info=True); return jsonify({'status': 'error', 'message': 'Internal server error.'}), 500

@app.route('/generate_pasted_comprehension', methods=['POST'])
def generate_pasted_comprehension_route():
    logging.info("Received request at /generate_pasted_comprehension")
    if client is None: return jsonify({'status': 'error', 'message': 'Server error: AI client not initialized'}), 500
    if not request.is_json: return jsonify({'status': 'error', 'message': 'Request must be JSON'}), 400
    try:
        data = request.get_json(); pasted_text = data.get('pasted_text'); num_questions_req = data.get('num_questions', 5)
        if not pasted_text or not pasted_text.strip(): return jsonify({'status': 'error', 'message': 'Missing "pasted_text"'}), 400
        try: num_questions = int(num_questions_req); assert 2 <= num_questions <= 15
        except (ValueError, TypeError, AssertionError): return jsonify({'status': 'error', 'message': 'Invalid num_questions (must be int 2-15)'}), 400
        logging.info(f"Pasted Text Comp request: NumQuestions={num_questions}, Text Length={len(pasted_text)}")
        prompt_content = create_pasted_comprehension_prompt(pasted_text, num_questions)
        message = client.messages.create(model=ANTHROPIC_MODEL_NAME, max_tokens=1500 + (num_questions * 100), temperature=0.7, messages=[{ "role": "user", "content": prompt_content }])
        if message.content and len(message.content) > 0 and hasattr(message.content[0], 'text'): generated_content = message.content[0].text
        else: raise ValueError("Failed to parse content from AI response.")
        logging.info(f"Generated Pasted Text Comp content length: {len(generated_content)} chars")
        return jsonify({'status': 'success', 'pasted_comprehension_content': generated_content.strip()})
    except ValueError as ve: logging.warning(f"Validation Error: {ve}"); return jsonify({'status': 'error', 'message': str(ve)}), 400
    except (anthropic.APIError, httpx.RequestError) as e: return handle_anthropic_error(e)
    except Exception as e: logging.error(f"Unexpected error in /generate_pasted_comprehension: {e}", exc_info=True); return jsonify({'status': 'error', 'message': 'Internal server error.'}), 500

@app.route('/generate_similar_questions', methods=['POST'])
def generate_similar_questions_route():
    logging.info("Received request at /generate_similar_questions")
    if client is None: return jsonify({'status': 'error', 'message': 'Server error: AI client not initialized'}), 500
    if not request.is_json: return jsonify({'status': 'error', 'message': 'Request must be JSON'}), 400
    try:
        data = request.get_json(); example_question = data.get('example_question'); num_questions_req = data.get('num_questions', 3); grade_level = data.get('grade_level', 'Not Specified')
        if not example_question or not example_question.strip(): return jsonify({'status': 'error', 'message': 'Missing "example_question"'}), 400
        try: num_questions = int(num_questions_req); assert 1 <= num_questions <= 10
        except (ValueError, TypeError, AssertionError): return jsonify({'status': 'error', 'message': 'Invalid num_questions (must be int 1-10)'}), 400
        logging.info(f"SimilarQ request: NumQuestions={num_questions}, Grade='{grade_level}', Example Length={len(example_question)}")
        prompt_content = create_similar_questions_prompt(example_question, num_questions, grade_level)
        message = client.messages.create(model=ANTHROPIC_MODEL_NAME, max_tokens=1000 + (num_questions * 200), temperature=0.75, messages=[{ "role": "user", "content": prompt_content }])
        if message.content and len(message.content) > 0 and hasattr(message.content[0], 'text'): generated_content = message.content[0].text
        else: raise ValueError("Failed to parse content from AI response.")
        logging.info(f"Generated SimilarQ content length: {len(generated_content)} chars")
        return jsonify({'status': 'success', 'similar_questions_content': generated_content.strip()})
    except ValueError as ve: logging.warning(f"Validation Error: {ve}"); return jsonify({'status': 'error', 'message': str(ve)}), 400
    except (anthropic.APIError, httpx.RequestError) as e: return handle_anthropic_error(e)
    except Exception as e: logging.error(f"Unexpected error in /generate_similar_questions: {e}", exc_info=True); return jsonify({'status': 'error', 'message': 'Internal server error.'}), 500

@app.route('/generate_worksheet', methods=['POST']) # This is the Gap Fill route
def generate_worksheet_route():
    logging.info("Received request at /generate_worksheet (Gap Fill)")
    if client is None: return jsonify({'status': 'error', 'message': 'Server error: AI client not initialized'}), 500
    if not request.is_json: return jsonify({'status': 'error', 'message': 'Request must be JSON'}), 400
    try:
        data = request.get_json(); topic = data.get('topic'); grade_level = data.get('grade_level', 'middle school'); num_sentences_req = data.get('num_sentences', 7)
        if not topic: return jsonify({'status': 'error', 'message': 'Missing "topic"'}), 400
        try: num_sentences = int(num_sentences_req); assert 3 <= num_sentences <= 15
        except (ValueError, TypeError, AssertionError): return jsonify({'status': 'error', 'message': 'Invalid num_sentences (must be int 3-15)'}), 400
        logging.info(f"Gap Fill request: Topic='{topic}', Grade='{grade_level}', NumSentences={num_sentences}")
        prompt_content = create_gap_fill_prompt(topic, grade_level, num_sentences)
        message = client.messages.create(model=ANTHROPIC_MODEL_NAME, max_tokens=500 + (num_sentences * 60), temperature=0.7, messages=[{ "role": "user", "content": prompt_content }])
        if message.content and len(message.content) > 0 and hasattr(message.content[0], 'text'): generated_content = message.content[0].text
        else: raise ValueError("Failed to parse content from AI response.")
        logging.info(f"Generated Gap Fill content length: {len(generated_content)} chars")
        return jsonify({'status': 'success', 'worksheet_content': generated_content.strip()})
    except ValueError as ve: logging.warning(f"Validation Error: {ve}"); return jsonify({'status': 'error', 'message': str(ve)}), 400
    except (anthropic.APIError, httpx.RequestError) as e: return handle_anthropic_error(e)
    except Exception as e: logging.error(f"Unexpected error in /generate_worksheet: {e}", exc_info=True); return jsonify({'status': 'error', 'message': 'Internal server error.'}), 500

# --- Word Search Routes ---
@app.route('/generate_word_list', methods=['POST'])
def generate_word_list_route():
    logging.info("Received request at /generate_word_list")
    if client is None: return jsonify({'status': 'error', 'message': 'Server error: AI client not initialized'}), 500
    if not request.is_json: return jsonify({'status': 'error', 'message': 'Request must be JSON'}), 400
    try:
        data = request.get_json(); topic = data.get('topic'); grade_level = data.get('grade_level', 'middle school'); num_words = int(data.get('num_words', 15))
        if not topic: return jsonify({'status': 'error', 'message': 'Missing "topic"'}), 400
        if not 5 <= num_words <= 30: return jsonify({'status': 'error', 'message': 'Number of words must be between 5 and 30'}), 400
        logging.info(f"Word List request: Topic='{topic}', Grade='{grade_level}', NumWords={num_words}")
        prompt_content = create_word_list_prompt(topic, grade_level, num_words)
        message = client.messages.create(model=ANTHROPIC_MODEL_NAME, max_tokens=300 + (num_words * 10), temperature=0.5, messages=[{ "role": "user", "content": prompt_content }])
        if message.content and len(message.content) > 0 and hasattr(message.content[0], 'text'): generated_content = message.content[0].text.strip()
        else: raise ValueError("Failed to parse content from AI response.")
        logging.info(f"Generated word list raw text length: {len(generated_content)} chars")
        return jsonify({'status': 'success', 'word_list_text': generated_content })
    except ValueError as ve: logging.error(f"Validation/Processing Error: {ve}", exc_info=True); return jsonify({'status': 'error', 'message': str(ve)}), 400
    except (anthropic.APIError, httpx.RequestError) as e: return handle_anthropic_error(e)
    except Exception as e: logging.error(f"Unexpected error in /generate_word_list: {e}", exc_info=True); return jsonify({'status': 'error', 'message': 'Internal server error.'}), 500

@app.route('/generate_word_search_grid', methods=['POST'])
def generate_word_search_grid_route():
    logging.info("Received request at /generate_word_search_grid")
    if not request.is_json: return jsonify({'status': 'error', 'message': 'Request must be JSON'}), 400
    try:
        data = request.get_json(); word_list_raw = data.get('word_list'); size_preference = data.get('size', 'medium'); allow_backwards = data.get('allow_backwards', False); mask = data.get('mask', None)
        if not word_list_raw or not isinstance(word_list_raw, list): return jsonify({'status': 'error', 'message': 'Missing or invalid "word_list"'}), 400
        size_map = {"small": 10, "medium": 13, "large": 15}; dimension = size_map.get(size_preference.lower(), 13)
        logging.info(f"Attempting grid generation: Size={dimension}x{dimension}, Backwards={allow_backwards}, Mask={mask}, Words={word_list_raw[:5]}...")
        grid_list, placed_words = generate_wordsearch_grid(wordlist=word_list_raw, nrows=dimension, ncols=dimension, allow_backwards_words=allow_backwards, mask_name=mask, attempts=15)
        html_output = '<div class="worksheet-section word-search-container">'; html_output += '<h3>Word Search</h3>'
        html_output += '<table class="word-search-grid" style="border-collapse: collapse; font-family: monospace; margin-bottom: 15px; border: 1px solid #ddd;">'
        for row_data in grid_list:
            html_output += '<tr>'
            for letter in row_data: display_letter = letter if letter.strip() else ' '; html_output += f'<td style="border: 1px solid #eee; width: 25px; height: 25px; text-align: center; vertical-align: middle; padding: 1px;">{display_letter}</td>'
            html_output += '</tr>'
        html_output += '</table>'; html_output += '<h3>Word List</h3>'
        html_output += '<ul class="word-search-list" style="list-style: none; padding-left: 0; columns: 2; -webkit-columns: 2; -moz-columns: 2;">'
        for word in sorted(placed_words): html_output += f'<li style="margin-bottom: 5px;">{word}</li>'
        html_output += '</ul></div>'; logging.info(f"Successfully generated word search grid HTML using internal logic.")
        return jsonify({'status': 'success', 'content_html': html_output})
    except (ValueError, PuzzleGenerationError) as gen_err: logging.warning(f"Word search generation failed: {gen_err}"); return jsonify({'status': 'error', 'message': str(gen_err)}), 400
    except Exception as e: logging.error(f"Unexpected error in /generate_word_search_grid: {e}", exc_info=True); return jsonify({'status': 'error', 'message': 'An internal server error occurred creating the word search.'}), 500

@app.route('/generate_keywords_definitions', methods=['POST'])
def generate_keywords_definitions_route():
    logging.info("Request received: /generate_keywords_definitions")
    if not request.is_json:
        logging.warning("Request aborted: Content-Type is not application/json.")
        return jsonify({'status': 'error', 'message': 'Request must be JSON'}), 400

    data = request.get_json()
    if not data:
        logging.warning("Request aborted: Empty JSON payload received.")
        return jsonify({'status': 'error', 'message': 'Empty JSON payload received.'}), 400

    topic = data.get('topic')
    grade_level = data.get('grade_level')
    try:
        # Default to 10 keywords if not provided or invalid
        num_keywords = int(data.get('num_keywords', 10))
        if num_keywords <= 0:
            num_keywords = 10
            logging.warning("Invalid num_keywords <= 0 received, defaulting to 10.")
    except (ValueError, TypeError):
        num_keywords = 10
        logging.warning("Invalid or missing num_keywords received, defaulting to 10.")


    if not topic or not grade_level:
        logging.warning("Request aborted: Missing 'topic' or 'grade_level' in JSON payload.")
        return jsonify({'status': 'error', 'message': 'Missing "topic" or "grade_level".'}), 400

    logging.debug(f"Parameters received: topic='{topic}', grade_level='{grade_level}', num_keywords={num_keywords}")

    # --- Construct the AI Prompt ---
    prompt = f"""Generate {num_keywords} unique keywords related to the topic "{topic}" suitable for grade level {grade_level}.
For each keyword, provide a concise and clear definition appropriate for that grade level.
Format the output STRICTLY as a JSON list of objects, where each object has a "keyword" key and a "definition" key.
Do not include any other text, explanations, or markdown formatting before or after the JSON list. Just the raw JSON.

Example format:
[
  {{"keyword": "Example Term 1", "definition": "This is the definition for term 1."}},
  {{"keyword": "Example Term 2", "definition": "This is the definition for term 2."}}
]

Begin JSON list now:"""

    try:
        logging.info(f"Sending request to Anthropic API for keywords/definitions...")
        message = client.messages.create(
            model=ANTHROPIC_MODEL_NAME, # Use your configured model name variable
            max_tokens=1500,  # Adjust as needed, maybe calculate based on num_keywords?
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )

        # Extract the text content from the response
        # Assuming the response structure gives text in message.content[0].text
        if not message.content or not message.content[0].text:
             raise ValueError("Received empty content from Anthropic API.")

        ai_response_text = message.content[0].text.strip()
        logging.debug(f"Raw AI response received:\n{ai_response_text}")

        # --- Parse the JSON Response ---
        # Sometimes the AI might still wrap the JSON in markdown backticks
        if ai_response_text.startswith("```json"):
            ai_response_text = ai_response_text[7:] # Remove ```json
        if ai_response_text.endswith("```"):
            ai_response_text = ai_response_text[:-3] # Remove ```
        ai_response_text = ai_response_text.strip() # Clean whitespace again

        try:
            keyword_data = json.loads(ai_response_text)
        except json.JSONDecodeError as json_err:
            logging.error(f"Failed to parse JSON from AI response: {json_err}")
            logging.error(f"Problematic AI response text: {ai_response_text}")
            # Consider trying to extract JSON manually if needed, but often better to fail
            return jsonify({'status': 'error', 'message': f'AI response format error (not valid JSON).'}), 500

        # --- Basic Validation of Parsed Data ---
        if not isinstance(keyword_data, list):
            logging.error(f"AI response parsed, but it's not a list. Type: {type(keyword_data)}")
            return jsonify({'status': 'error', 'message': 'AI response format error (expected a list).'}), 500

        validated_data = []
        for item in keyword_data:
            if isinstance(item, dict) and 'keyword' in item and 'definition' in item:
                # Basic cleaning
                kw = str(item['keyword']).strip()
                df = str(item['definition']).strip()
                if kw and df: # Ensure they are not empty after stripping
                     validated_data.append({'keyword': kw, 'definition': df})
            else:
                logging.warning(f"Skipping invalid item in AI response list: {item}")

        if not validated_data:
             logging.error("AI response parsed as list, but no valid keyword/definition pairs found.")
             return jsonify({'status': 'error', 'message': 'AI returned no valid keyword/definition pairs.'}), 500

        logging.info(f"Successfully generated and parsed {len(validated_data)} keywords/definitions.")
        return jsonify({'status': 'success', 'keywords': validated_data})

    except Exception as e:
        # Catch errors from the API call itself or other unexpected issues
        logging.error(f"Error during keywords/definitions generation: {e}", exc_info=True)
        # Provide a generic error to the frontend
        return jsonify({'status': 'error', 'message': 'An internal error occurred while generating keywords.'}), 500

# --- Item/Worksheet Persistence Routes ---
@app.route('/save_item', methods=['POST'])
def save_item_route():
    # --- 1. Log that the function started ---
    logging.info("--- save_item route entered ---")

    # --- 2. Check if the incoming request is valid JSON ---
    if not request.is_json:
        logging.error("save_item: Request is not JSON")
        return jsonify({'status': 'error', 'message': 'Request must be JSON'}), 400

    # --- 3. Get the data sent from the frontend ---
    data = request.get_json()
    logging.debug(f"save_item: Received data keys: {list(data.keys())}")

    # --- 4. Extract all the pieces of data we expect ---
    item_type = data.get('item_type')
    source_topic = data.get('source_topic')
    source_url = data.get('source_url')
    grade_level = data.get('grade_level')
    item_data_json = data.get('item_data_json', None) # Get the extra JSON data (like image boxes), default to None if missing
    content_html = data.get('content_html')
    existing_item_id = data.get('existing_item_id') # <<<--- Check if the frontend sent an ID for an item it wants to UPDATE

    # --- 5. Make sure we have the essential pieces ---
    if not item_type or not grade_level or not content_html:
        logging.error(f"save_item: Missing required fields. Got type: {item_type}, grade: {grade_level}, content: {bool(content_html)}")
        return jsonify({'status': 'error', 'message': 'Missing required fields: item_type, grade_level, content_html'}), 400

    # --- 6. Prepare variables ---
    item_to_save = None # This will hold the database item (either found or newly created)
    is_update = False   # This flag tells us if we are updating or creating

    # --- 7. Try to handle database actions (finding, updating, creating) ---
    try:
        # --- 8. Check if an ID for updating was provided ---
        if existing_item_id:
            logging.debug(f"save_item: Received potential existing_item_id: {existing_item_id}")
            try:
                # --- 9. Try to convert the ID to a number and find the item in the database ---
                item_id_int = int(existing_item_id)
                item_to_save = db.session.get(GeneratedItem, item_id_int) # Use db.session.get to find by primary key

                # --- 10. If we found the item, set the flag to true ---
                if item_to_save:
                    is_update = True
                    logging.info(f"save_item: Found existing item ID {item_id_int} for update.")
                else:
                    # If ID was sent but not found, log a warning and proceed to create a NEW item instead.
                    logging.warning(f"save_item: Received existing_item_id {item_id_int}, but item not found in DB. Will create new.")
            except (ValueError, TypeError):
                 # If the ID sent wasn't a valid number, log a warning and proceed to create a NEW item.
                 logging.warning(f"save_item: Invalid existing_item_id format received: {existing_item_id}. Will create new.")
                 # Let the code fall through to the 'create new' section below

        # --- 11. If the 'is_update' flag is true, update the found item ---
        if is_update and item_to_save: # Double-check item_to_save exists
            logging.debug(f"save_item: Updating item ID {item_to_save.id}")
            # --- 12. Update all the fields of the found item with the new data ---
            item_to_save.item_type = item_type
            item_to_save.source_topic = source_topic
            item_to_save.source_url = source_url
            item_to_save.grade_level = grade_level
            item_to_save.item_data_json = item_data_json # Update the extra JSON data too
            item_to_save.content_html = content_html
            # The 'last_modified_date' will update automatically because of 'onupdate' in the model definition
            logging.debug(f"save_item: Item object UPDATED in memory (before commit): {item_to_save}")

        # --- 13. If the 'is_update' flag is false, it means we need to create a new item ---
        else:
            logging.debug(f"save_item: Preparing to create NEW item: type={item_type}, grade={grade_level}")
            # --- 14. Create a new GeneratedItem object with all the data ---
            item_to_save = GeneratedItem(
                item_type=item_type,
                source_topic=source_topic,
                source_url=source_url,
                grade_level=grade_level,
                item_data_json=item_data_json, # Store the extra JSON data
                content_html=content_html
                # 'creation_date' and 'last_modified_date' will get default values automatically
            )
            logging.debug(f"save_item: GeneratedItem object CREATED: {item_to_save}")
            # --- 15. Add this new item object to the database session ---
            db.session.add(item_to_save)
            logging.debug("save_item: Added NEW item to session.")

        # --- 16. Commit the changes to the database (this saves the update OR the new item) ---
        db.session.commit()
        action = "updated" if is_update else "created" # Figure out which word to use in the log message
        logging.info(f"save_item: Commit successful. Item {action}. ID assigned/confirmed: {item_to_save.id}")

        # --- 17. Send a success response back to the frontend, including the item's ID ---
        response_data = {'status': 'success', 'item_id': item_to_save.id}
        logging.debug(f"save_item: Returning success data: {response_data}")
        return jsonify(response_data)

    # --- 18. If anything went wrong in the 'try' block above... ---
    except Exception as e:
        db.session.rollback() # Undo any changes made in this session before the error
        logging.error(f"save_item: Exception during DB operation: {e}", exc_info=True) # Log the full error
        # --- 19. Send an error response back to the frontend ---
        return jsonify({'status': 'error', 'message': f'Database error: {str(e)}'}), 500
@app.route('/get_item/<int:item_id>', methods=['GET'])
def get_item_route(item_id):
    logging.info(f"Received request at /get_item/{item_id}")
    try:
        item = db.session.get(GeneratedItem, item_id)
        if item is None: logging.warning(f"Item with ID {item_id} not found."); return jsonify({'status': 'error', 'message': 'Item not found in library.'}), 404
        logging.info(f"Found item: {item.item_type}")
        item_data = {'id': item.id, 'item_type': item.item_type, 'source_topic': item.source_topic, 'source_url': item.source_url,
                       'grade_level': item.grade_level, 'item_data_json': item.item_data_json, 'content_html': item.content_html, 'creation_date': item.creation_date.isoformat(),
                       'last_modified_date': item.last_modified_date.isoformat()}
        return jsonify({'status': 'success', 'item': item_data})
    except Exception as e: logging.error(f"Error retrieving item {item_id} from database: {e}", exc_info=True); return jsonify({'status': 'error', 'message': 'Error retrieving item details.'}), 500

@app.route('/save_worksheet', methods=['POST'])
def save_worksheet_route():
    logging.info("Received request at /save_worksheet")
    if not request.is_json: return jsonify({'status': 'error', 'message': 'Request must be JSON'}), 400
    data = request.get_json(); worksheet_title = data.get('title', 'Untitled Worksheet'); item_ids_ordered = data.get('item_ids')
    if not item_ids_ordered or not isinstance(item_ids_ordered, list): return jsonify({'status': 'error', 'message': 'Missing or invalid "item_ids" list'}), 400
    logging.info(f"Attempting to save worksheet: Title='{worksheet_title}', Item IDs={item_ids_ordered}")
    try:
        new_worksheet = Worksheet(title=worksheet_title); db.session.add(new_worksheet); db.session.flush(); worksheet_id = new_worksheet.id
        items_to_add = []; order_index = 0
        for item_id in item_ids_ordered:
            generated_item = db.session.get(GeneratedItem, int(item_id))
            if generated_item:
                assoc_insert = worksheet_items_association.insert().values(worksheet_id=worksheet_id, generated_item_id=generated_item.id, item_order=order_index)
                db.session.execute(assoc_insert); items_to_add.append(generated_item.id); order_index += 1
            else: logging.warning(f"GeneratedItem ID {item_id} not found while saving worksheet {worksheet_id}. Skipping.")
        db.session.commit(); logging.info(f"Successfully saved Worksheet ID: {worksheet_id} with item IDs: {items_to_add} in order.")
        return jsonify({'status': 'success', 'message': 'Worksheet saved successfully!', 'worksheet_id': worksheet_id})
    except Exception as e: db.session.rollback(); logging.error(f"Error saving worksheet to database: {e}", exc_info=True); return jsonify({'status': 'error', 'message': f'Database error while saving worksheet: {e}'}), 500

@app.route('/list_worksheets', methods=['GET'])
def list_worksheets_route():
    logging.info("Received request at /list_worksheets")
    try:
        worksheets = Worksheet.query.order_by(Worksheet.last_modified_date.desc()).limit(50).all()
        worksheets_list = [{'id': ws.id, 'title': ws.title, 'creation_date': ws.creation_date.isoformat(),
                            'last_modified_date': ws.last_modified_date.isoformat()} for ws in worksheets]
        return jsonify({'status': 'success', 'worksheets': worksheets_list})
    except Exception as e: logging.error(f"Error retrieving worksheets from database: {e}", exc_info=True); return jsonify({'status': 'error', 'message': 'Error retrieving worksheet list.'}), 500

@app.route('/load_worksheet/<int:worksheet_id>', methods=['GET'])
def load_worksheet_route(worksheet_id):
    logging.info(f"Received request at /load_worksheet/{worksheet_id}")
    try:
        worksheet = db.session.get(Worksheet, worksheet_id)
        if worksheet is None: return jsonify({'status': 'error', 'message': 'Worksheet not found.'}), 404
        ordered_items = worksheet.items # Relationship handles ordering
        items_data = []
        for item in ordered_items:
            items_data.append({
                'id': item.id,
                'item_type': item.item_type,
                'source_topic': item.source_topic,
                'source_url': item.source_url,
                'grade_level': item.grade_level,
                'content_html': item.content_html,
                'item_data_json': item.item_data_json # <<< ADD THIS LINE
            })
        logging.info(f"Returning {len(items_data)} items for Worksheet ID: {worksheet_id}")
        return jsonify({'status': 'success', 'worksheet_title': worksheet.title, 'items': items_data})
    except Exception as e: logging.error(f"Error retrieving worksheet {worksheet_id} from database: {e}", exc_info=True); return jsonify({'status': 'error', 'message': 'Error retrieving worksheet details.'}), 500
@app.route('/generate_similar_written_questions', methods=['POST'])
def generate_similar_written_questions_route():
    logging.info("Request received: /generate_similar_written_questions")
    if not request.is_json:
        logging.warning("Request aborted: Content-Type is not application/json.")
        return jsonify({'status': 'error', 'message': 'Request must be JSON'}), 400

    data = request.get_json()
    if not data:
        logging.warning("Request aborted: Empty JSON payload received.")
        return jsonify({'status': 'error', 'message': 'Empty JSON payload received.'}), 400

    example_question = data.get('example_question')
    grade_level = data.get('grade_level') # Might be "Not Specified"
    try:
        num_questions = int(data.get('num_questions', 3))
        if num_questions <= 0:
            num_questions = 3
            logging.warning("Invalid num_questions <= 0 received, defaulting to 3.")
    except (ValueError, TypeError):
        num_questions = 3
        logging.warning("Invalid or missing num_questions received, defaulting to 3.")

    if not example_question:
        logging.warning("Request aborted: Missing 'example_question' in JSON payload.")
        return jsonify({'status': 'error', 'message': 'Missing "example_question".'}), 400

    # Use grade level if provided, otherwise let AI infer
    grade_level_text = f"potentially intended for grade level '{grade_level}'" if grade_level and grade_level != "Not Specified" else "with an unspecified grade level"

    logging.debug(f"Parameters: num_questions={num_questions}, grade_level='{grade_level}', example='{example_question[:100]}...'")

    # --- Construct the AI Prompt ---
    prompt = f"""You are an expert assessment designer tasked with creating question variations.
Analyze the following example question, {grade_level_text}:

--- Example Question ---
{example_question}
--- End Example Question ---

First, identify and concisely describe the primary skill(s) being assessed by this question.
Second, generate {num_questions} new, distinct questions that assess the *exact same skill(s)* you identified. Ensure these new questions are appropriate for the specified grade level (or maintain the implied level of the example), keep a similar difficulty, but use different scenarios, contexts, or specific details. Avoid simple rephrasing of the example.

Format your response strictly as a JSON object containing two keys:
1. "skills_identified": A string describing the core skill(s) tested (keep this brief, maybe 1-2 sentences).
2. "similar_questions": A JSON list containing exactly {num_questions} strings, where each string is one of the new questions.

Do not include any explanations or text outside the JSON object.

Example JSON Output Format:
{{
  "skills_identified": "Applying the definition of photosynthesis to identify necessary components.",
  "similar_questions": [
    "Besides sunlight, what is one essential reactant that plants need for photosynthesis?",
    "Which part of a plant cell is primarily responsible for carrying out photosynthesis?",
    "If a plant is kept in complete darkness but given water and carbon dioxide, can it perform photosynthesis? Explain why or why not."
  ]
}}

Begin JSON object now:"""

    try:
        logging.info(f"Sending request to Anthropic API for similar written questions...")
        message = client.messages.create(
            model=ANTHROPIC_MODEL_NAME,
            max_tokens=1500, # Adjust as needed
            messages=[{"role": "user", "content": prompt}]
        )

        if not message.content or not message.content[0].text:
             raise ValueError("Received empty content from Anthropic API.")

        ai_response_text = message.content[0].text.strip()
        logging.debug(f"Raw AI response received:\n{ai_response_text}")

        # Clean potential markdown backticks
        if ai_response_text.startswith("```json"):
            ai_response_text = ai_response_text[7:]
        if ai_response_text.endswith("```"):
            ai_response_text = ai_response_text[:-3]
        ai_response_text = ai_response_text.strip()

        try:
            parsed_data = json.loads(ai_response_text)
        except json.JSONDecodeError as json_err:
            logging.error(f"Failed to parse JSON from AI response: {json_err}")
            logging.error(f"Problematic AI response text for similar written q: {ai_response_text}")
            return jsonify({'status': 'error', 'message': 'AI response format error (not valid JSON).'}), 500

        # Validate structure
        if not isinstance(parsed_data, dict):
             logging.error(f"Parsed data is not a dict: {type(parsed_data)}")
             return jsonify({'status': 'error', 'message': 'AI response format error (expected JSON object).'}), 500

        skills = parsed_data.get("skills_identified")
        questions = parsed_data.get("similar_questions")

        if not isinstance(skills, str) or not isinstance(questions, list):
             logging.error(f"Missing/invalid keys in parsed JSON. Skills type: {type(skills)}, Questions type: {type(questions)}")
             return jsonify({'status': 'error', 'message': 'AI response format error (missing/invalid keys).'}), 500

        # Optional: Validate number of questions returned matches request?
        # if len(questions) != num_questions:
        #    logging.warning(f"AI returned {len(questions)} questions, expected {num_questions}")

        validated_questions = [str(q).strip() for q in questions if isinstance(q, str) and str(q).strip()]

        logging.info(f"Successfully generated {len(validated_questions)} similar written questions.")
        # Return the whole parsed object nested under 'data'
        return jsonify({
            'status': 'success',
            'data': {
                 'skills_identified': skills.strip(),
                 'similar_questions': validated_questions
                 }
            })

    except Exception as e:
        logging.error(f"Error during similar written questions generation: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': 'An internal error occurred while generating questions.'}), 500
# --- Export and Frontend Routes ---
@app.route('/export/docx/<int:worksheet_id>')
def export_worksheet_docx(worksheet_id):
    """Exports a specific worksheet as a DOCX file."""
    logging.info(f"Request: /export/docx/{worksheet_id}")
    try:
        worksheet = Worksheet.query.get_or_404(worksheet_id)
        items = worksheet.items # Assumes items are correctly ordered by the relationship
        logging.info(f"Found worksheet '{worksheet.title}' with {len(items)} items.")

        document = docx.Document()
        logging.debug("Initialized docx Document.")

        # --- Apply Document Formatting ---
        target_font='Century Gothic' # Or 'Arial', 'Calibri' etc.
        try:
            normal_style = document.styles['Normal']
            normal_style.font.name = target_font
            normal_style.paragraph_format.space_after = Pt(6)
            logging.info(f"Set Normal font to {target_font} and paragraph spacing.")
        except Exception as font_e:
            logging.warning(f"Could not set default font/style: {font_e}")

        try:
            h3_style = document.styles['Heading 3']
            h3_style.font.name = target_font
            h3_style.paragraph_format.space_before = Pt(12)
            h3_style.paragraph_format.space_after = Pt(6)
            logging.info(f"Set H3 font to {target_font} and spacing.")
        except Exception as style_e:
            logging.warning(f"Could not modify Heading 3 style: {style_e}")

        try: # Set Margins
            section = document.sections[0]
            section.left_margin=Inches(0.75); section.right_margin=Inches(0.75) # Slightly wider than 0.5
            section.top_margin=Inches(0.75); section.bottom_margin=Inches(0.75)
            logging.info("Set document margins to 0.75 inches.")
        except Exception as margin_e:
            logging.warning(f"Could not set margins: {margin_e}")

        document.add_heading(worksheet.title, level=1)
        document.add_paragraph()
        logging.debug(f"Added worksheet title: {worksheet.title}")

        logging.info("Starting item processing loop...")
        for item_index, item in enumerate(items):
            logging.debug(f"--- Processing item index: {item_index}, ID: {item.id}, Type: {item.item_type} ---")
            soup = BeautifulSoup(item.content_html, 'lxml')
            processed_as_type = "Unknown"

            word_search_container = soup.find('div', class_='word-search-container')

            if word_search_container:
                processed_as_type = "WordSearch"
                logging.debug("Item identified as Word Search.")
                try:
                    heading_ws_el = word_search_container.find('h3')
                    document.add_heading(heading_ws_el.get_text(strip=True) if heading_ws_el else "Word Search", level=3)

                    grid_table = word_search_container.find('table', class_='word-search-grid')
                    grid_list = []
                    if grid_table:
                        for row in grid_table.find_all('tr'):
                            grid_list.append([cell.get_text(strip=True) for cell in row.find_all('td')])
                    
                    if grid_list:
                        grid_image = create_wordsearch_image(grid_list)
                        if grid_image:
                            img_buffer = io.BytesIO(); grid_image.save(img_buffer, format='PNG'); img_buffer.seek(0)
                            try: document.add_picture(img_buffer, width=Inches(6.0)) # Adjust width as needed
                            except Exception as pic_e: logging.error(f"Failed to add picture: {pic_e}"); document.add_paragraph("[Error adding grid image]", style='Comment')
                        else: logging.warning("create_wordsearch_image returned None."); document.add_paragraph("[Error generating grid image]", style='Comment')
                    else: logging.warning("Word search grid table not found in HTML."); document.add_paragraph("[Word search grid table not found in HTML]", style='Comment')

                    word_list_headings = word_search_container.find_all('h3')
                    word_list_heading_text = "Word List"
                    if len(word_list_headings) > 1: # Assumes second H3 is for word list
                        word_list_heading_text = word_list_headings[1].get_text(strip=True)
                    document.add_heading(word_list_heading_text, level=3)
                    
                    word_list_ul = word_search_container.find('ul', class_='word-search-list')
                    words_to_find = [li.get_text(strip=True) for li in word_list_ul.find_all('li')] if word_list_ul else []
                    words_to_find = [word for word in words_to_find if word]

                    if words_to_find:
                        # Simple paragraph for words, consider columns/table for better formatting later
                        p_words = document.add_paragraph()
                        for i, word in enumerate(words_to_find):
                            p_words.add_run(word)
                            if i < len(words_to_find) - 1:
                                p_words.add_run("     ") # Add spacing between words
                    else: document.add_paragraph("[Word list empty or not found in HTML]", style='Comment')
                    document.add_paragraph()
                except Exception as ws_err:
                    logging.error(f"Error processing WordSearch item {item.id}: {ws_err}", exc_info=True)
                    document.add_paragraph(f"[Error processing Word Search item {item.id}]", style='Comment')

            elif item.item_type.startswith('keywords-'):
                processed_as_type = "KeywordsTable"
                logging.debug(f"Processing keywords item {item.id} (Type: {item.item_type}) for DOCX.")
                try:
                    html_table = soup.find('table', class_='keywords-table') # soup is already defined
                    if not html_table: raise ValueError("Could not find 'table.keywords-table' in HTML")
                    rows = html_table.find_all('tr')
                    if not rows: raise ValueError("Keywords table found but contained no rows.")
                    header_cells = rows[0].find_all(['th', 'td'])
                    num_cols = len(header_cells)
                    if num_cols <= 0: raise ValueError("Keywords table found, but header row had no columns.")
                    data_rows = rows[1:]
                    num_data_rows = len(data_rows)

                    docx_table = document.add_table(rows=num_data_rows + 1, cols=num_cols)
                    docx_table.style = 'Table Grid'

                    for j, cell in enumerate(header_cells):
                        if j < num_cols:
                            header_text = ' '.join(cell.get_text(strip=True).split())
                            hdr_cell = docx_table.cell(0, j)
                            hdr_cell.text = header_text
                            if hdr_cell.paragraphs and hdr_cell.paragraphs[0].runs: hdr_cell.paragraphs[0].runs[0].font.bold = True
                            elif hdr_cell.paragraphs: hdr_cell.paragraphs[0].add_run().font.bold = True
                    
                    for i, html_row in enumerate(data_rows):
                        cells = html_row.find_all('td')
                        for j, cell in enumerate(cells):
                            if j < num_cols:
                                cell_text = ' '.join(cell.get_text(strip=True).split())
                                docx_table.cell(i + 1, j).text = cell_text if cell_text else ""
                    
                    if item.item_type == 'keywords-matching' and num_cols == 3:
                        middle_col_index = 1
                        total_rows_in_table = num_data_rows + 1
                        for i_row in range(total_rows_in_table):
                            try:
                                cell_to_modify = docx_table.cell(i_row, middle_col_index)
                                tcPr = cell_to_modify._tc.get_or_add_tcPr()
                                tcBorders = tcPr.first_child_found_in("w:tcBorders")
                                if tcBorders is None: tcBorders = OxmlElement("w:tcBorders"); tcPr.append(tcBorders)
                                for border_name in ['top', 'bottom']: # Only remove top and bottom
                                    border_el = OxmlElement(f'w:{border_name}'); border_el.set(qn('w:val'), 'nil'); tcBorders.append(border_el)
                            except Exception as border_err: logging.error(f"Error modifying borders cell({i_row},{middle_col_index}): {border_err}")
                    document.add_paragraph()
                except Exception as parse_err:
                    logging.error(f"Error processing keywords table item {item.id}: {parse_err}", exc_info=True)
                    document.add_paragraph(f"[Error processing keyword table for item {item.id}]", style='Comment')

            elif item.item_type == 'similarWrittenQ':
                processed_as_type = "SimilarWrittenQ"
                logging.debug(f"Processing similarWrittenQ item {item.id} for DOCX.")
                skills = "Skills not specified."
                questions = []
                try:
                    if item.item_data_json:
                        parsed_data = json.loads(item.item_data_json)
                        if isinstance(parsed_data, dict):
                            skills = parsed_data.get('skills_identified', skills).strip()
                            questions_raw = parsed_data.get('similar_questions', [])
                            if isinstance(questions_raw, list):
                                questions = [str(q).strip() for q in questions_raw if str(q).strip()]
                    else: logging.warning(f"item_data_json not found for similarWrittenQ item {item.id}.")

                    document.add_heading("Skills Tested", level=3)
                    document.add_paragraph(skills)
                    document.add_paragraph()
                    document.add_heading("Generated Questions", level=3)
                    if questions:
                        for q_text in questions: document.add_paragraph(q_text, style='List Number')
                        p_reset = document.add_paragraph(); p_reset.add_run().font.size = Pt(1)
                    else: document.add_paragraph("[No similar questions found in saved data]", style='Comment')
                    document.add_paragraph()
                except Exception as export_err:
                    logging.error(f"Error adding similarWrittenQ item {item.id} to DOCX: {export_err}", exc_info=True)
                    document.add_paragraph(f"[Error processing Similar Written Questions item {item.id}]", style='Comment')
            
            elif item.item_type == 'imageLabel':
                processed_as_type = "ImageLabel"
                logging.debug(f"DOCX Export: Processing imageLabel item ID: {item.id}")
                try:
                    if not item.item_data_json: raise ValueError(f"Missing item_data_json for imageLabel item {item.id}")
                    data = json.loads(item.item_data_json)
                    if not isinstance(data, dict) or 'baseImage' not in data: raise ValueError(f"Invalid structure in item_data_json for item {item.id}")
                    base_image_filename = data['baseImage']
                    boxes_to_process = data.get('boxes_pct', data.get('boxes')) # Prefer _pct
                    is_percent_data = 'boxes_pct' in data
                    
                    image_path = os.path.join(app.config['UPLOAD_FOLDER'], base_image_filename)
                    if not os.path.exists(image_path): raise FileNotFoundError(f"Base image file not found: {image_path}")

                    img = Image.open(image_path).convert("RGB")
                    draw = ImageDraw.Draw(img)
                    
                    calc_ref_width_for_pct = None
                    calc_ref_height_for_pct = None

                    if is_percent_data: # This 'is_percent_data' was defined earlier based on 'boxes_pct' key presence
                        calc_ref_width_for_pct = data.get('refEditorWidth')
                        calc_ref_height_for_pct = data.get('refEditorHeight')
                        if not calc_ref_width_for_pct or not calc_ref_height_for_pct:
                            logging.warning(f"DOCX Export: Item {item.id} uses percentages (boxes_pct) but refEditorWidth/Height missing from JSON. "
                                            f"Falling back to naturalWidth/Height if available, then actual image dimensions. "
                                            f"This might lead to scaling issues if the original editor view was different from natural size.")
                            calc_ref_width_for_pct = data.get('naturalWidth', img.width)
                            calc_ref_height_for_pct = data.get('naturalHeight', img.height)
                        logging.debug(f"DOCX Export: For item {item.id} (percent data), using reference for % calc: {calc_ref_width_for_pct}x{calc_ref_height_for_pct}. Drawing on image of size: {img.width}x{img.height}")
                    else: # Old pixel data
                        # For old pixel data, we assume these pixels were meant for the natural image size,
                        # or if refEditorWidth was stored, maybe that. This path is less certain for old data.
                        # This will be the reference against which the old pixel data is scaled to the current image.
                        calc_ref_width_for_pixels = data.get('refEditorWidth', data.get('naturalWidth', img.width))
                        calc_ref_height_for_pixels = data.get('refEditorHeight', data.get('naturalHeight', img.height))
                        logging.debug(f"DOCX Export: For item {item.id} (pixel data), reference dimensions for old pixels: {calc_ref_width_for_pixels}x{calc_ref_height_for_pixels}. Drawing on image of size: {img.width}x{img.height}")

                    font_main, font_small = None, None
                    try:
                        font_main = ImageFont.truetype("arial.ttf", 14) # Adjust font size as needed
                        font_small = ImageFont.truetype("arial.ttf", 10)
                    except IOError: 
                        logging.warning("Arial font not found, using default for image labels in DOCX.")
                        try:
                            font_main = ImageFont.load_default()
                            font_small = ImageFont.load_default() # Use same default if specific size fails
                        except Exception as font_load_err:
                            logging.error(f"Could not load default PIL font: {font_load_err}")
                            # font_main and font_small will remain None
                    
                    if boxes_to_process and isinstance(boxes_to_process, list):
                        for box_data in boxes_to_process:
                            try:
                                px_x, px_y, px_w, px_h = 0.0, 0.0, 0.0, 0.0 # Pixels to draw on the final 'img'

                                if is_percent_data:
                                    # Convert percentages to absolute pixel values relative to calc_ref_width_for_pct
                                    x_abs_on_ref = (float(box_data.get('x_pct', 0)) / 100.0) * calc_ref_width_for_pct
                                    y_abs_on_ref = (float(box_data.get('y_pct', 0)) / 100.0) * calc_ref_height_for_pct
                                    w_abs_on_ref = (float(box_data.get('width_pct', 10)) / 100.0) * calc_ref_width_for_pct # Default 10%
                                    h_abs_on_ref = (float(box_data.get('height_pct', 5)) / 100.0) * calc_ref_height_for_pct  # Default 5%
                                    
                                    # Now, these absolute pixel values (calculated on the reference dimensions)
                                    # need to be scaled if calc_ref_width_for_pct is different from the actual img.width
                                    # (i.e., if the editor view was scaled relative to the natural image when percentages were saved).
                                    # We are drawing onto 'img' (the natural-sized image).
                                    if calc_ref_width_for_pct == 0 or calc_ref_height_for_pct == 0: # Avoid division by zero
                                        logging.error(f"DOCX Export: calc_ref_width_for_pct or calc_ref_height_for_pct is zero for item {item.id}. Skipping box scaling.")
                                        px_x, px_y, px_w, px_h = x_abs_on_ref, y_abs_on_ref, w_abs_on_ref, h_abs_on_ref # Use unscaled, might be wrong
                                    else:
                                        scale_factor_to_natural_w = img.width / calc_ref_width_for_pct
                                        scale_factor_to_natural_h = img.height / calc_ref_height_for_pct

                                        px_x = x_abs_on_ref * scale_factor_to_natural_w
                                        px_y = y_abs_on_ref * scale_factor_to_natural_h
                                        px_w = w_abs_on_ref * scale_factor_to_natural_w
                                        px_h = h_abs_on_ref * scale_factor_to_natural_h
                                else: # Old pixel data (is_percent_data is False)
                                    # Assume old pixel data was meant to be drawn directly onto an image of 'calc_ref_width_for_pixels'
                                    # And then scale it to the current natural image size ('img.width', 'img.height').
                                    x_old_px = float(box_data.get('x', 0))
                                    y_old_px = float(box_data.get('y', 0))
                                    w_old_px = float(box_data.get('width', 50))  # Default 50px
                                    h_old_px = float(box_data.get('height', 30)) # Default 30px

                                    if calc_ref_width_for_pixels == 0 or calc_ref_height_for_pixels == 0:
                                        logging.error(f"DOCX Export: calc_ref_width_for_pixels or calc_ref_height_for_pixels is zero for item {item.id}. Skipping box scaling for old pixel data.")
                                        px_x, px_y, px_w, px_h = x_old_px, y_old_px, w_old_px, h_old_px
                                    else:
                                        scale_factor_to_natural_w = img.width / calc_ref_width_for_pixels
                                        scale_factor_to_natural_h = img.height / calc_ref_height_for_pixels

                                        px_x = x_old_px * scale_factor_to_natural_w
                                        px_y = y_old_px * scale_factor_to_natural_h
                                        px_w = w_old_px * scale_factor_to_natural_w
                                        px_h = h_old_px * scale_factor_to_natural_h
                                
                                # Draw semi-transparent box
                                # Ensure width and height are at least 1 pixel for drawing
                                draw_w = max(1, int(math.ceil(px_w)))
                                draw_h = max(1, int(math.ceil(px_h)))
                                rect_img_surface = Image.new('RGBA', (draw_w, draw_h), (255, 255, 255, 0)) # Transparent base
                                rect_draw_on_surface = ImageDraw.Draw(rect_img_surface)
                                rect_draw_on_surface.rectangle([(0,0), (draw_w-1, draw_h-1)], outline="black", width=1, fill=(255,255,255,255)) # Semi-transparent white
                                img.paste(rect_img_surface, (int(px_x), int(px_y)), rect_img_surface) # Paste with alpha blending

                                text_content = str(box_data.get('text', ''))
                                if box_data.get('type') != 'blank' and text_content:
                                    font_to_use = font_small if px_h < 25 and font_small else font_main 
                                    if font_to_use:
                                        try: 
                                            center_x = px_x + px_w / 2
                                            center_y = px_y + px_h / 2
                                            draw.text((center_x, center_y), text_content, fill="black", font=font_to_use, anchor="mm")
                                        except (AttributeError, TypeError): 
                                            text_bbox_fallback = draw.textbbox((0,0), text_content, font=font_to_use)
                                            text_w_fallback = text_bbox_fallback[2] - text_bbox_fallback[0]
                                            text_h_fallback = text_bbox_fallback[3] - text_bbox_fallback[1]
                                            ascent_fb, descent_fb = font_to_use.getmetrics()
                                            actual_text_visual_height_fb = ascent_fb 
                                            
                                            draw_text_x_fallback = px_x + (px_w - text_w_fallback) / 2
                                            draw_text_y_fallback = px_y + (px_h - actual_text_visual_height_fb) / 2 
                                            if descent_fb > 0 : # Basic adjustment for descenders if present
                                                draw_text_y_fallback -= descent_fb / 2 
                                            draw.text((draw_text_x_fallback, draw_text_y_fallback), text_content, fill="black", font=font_to_use)
                            except Exception as box_draw_err: 
                                logging.error(f"Error drawing a specific box/text for imageLabel in DOCX: {box_draw_err}", exc_info=True)
                    
                    img_buffer = io.BytesIO(); img.save(img_buffer, format='PNG'); img_buffer.seek(0)
                    page_width_inches = document.sections[0].page_width.inches - document.sections[0].left_margin.inches - document.sections[0].right_margin.inches
                    document.add_picture(img_buffer, width=Inches(min(6.5, page_width_inches)))
                    document.add_paragraph()
                except Exception as img_label_err:
                    logging.error(f"DOCX Export: Unexpected error processing imageLabel item {item.id}: {img_label_err}", exc_info=True)
                    document.add_paragraph(f"[Unexpected Server Error processing Image Labelling item {item.id}. Check server logs.]", style='Comment')
            
            else: # Generic handler for other text-based items
                processed_as_type = "GenericText"
                logging.debug(f"Processing generic text-based item {item.id} (Type: {item.item_type}) for DOCX.")
                try:
                    content_container = soup.find('body') # Or a more specific container if universally present
                    if content_container:
                        for element in content_container.children:
                            if not hasattr(element, 'name') or element.name is None:
                                if isinstance(element, str) and element.strip():
                                    document.add_paragraph(element.strip())
                                continue

                            if element.name == 'h3':
                                document.add_heading(element.get_text(strip=True), level=3)
                            elif element.name == 'p':
                                para = document.add_paragraph()
                                add_runs_from_html_element(para, element)
                            elif element.name in ['ol', 'ul']:
                                list_style = 'ListNumber' if element.name == 'ol' else 'ListBullet'
                                for li in element.find_all('li', recursive=False):
                                    para = document.add_paragraph(style=list_style)
                                    add_runs_from_html_element(para, li)
                                p_reset = document.add_paragraph(); p_reset.add_run().font.size = Pt(1) # Reset list numbering
                            elif element.name == 'div' and 'worksheet-section' in element.get('class', []):
                                section_title_el = element.find('h3', recursive=False)
                                if section_title_el:
                                    document.add_heading(section_title_el.get_text(strip=True), level=3)
                                    section_title_el.decompose() # Avoid re-processing

                                for child_el in element.children: # Process children of worksheet-section
                                    if not hasattr(child_el, 'name') or child_el.name is None:
                                        if isinstance(child_el, str) and child_el.strip(): document.add_paragraph(child_el.strip())
                                        continue
                                    if child_el.name == 'p':
                                        para = document.add_paragraph(); add_runs_from_html_element(para, child_el)
                                    elif child_el.name in ['ol', 'ul']:
                                        child_list_style = 'ListNumber' if child_el.name == 'ol' else 'ListBullet'
                                        for li_child in child_el.find_all('li', recursive=False):
                                            para = document.add_paragraph(style=child_list_style); add_runs_from_html_element(para, li_child)
                                        p_reset_child = document.add_paragraph(); p_reset_child.add_run().font.size = Pt(1)
                                    elif child_el.name == 'div': # E.g., answer key divs
                                        # Check for specific IDs for more targeted formatting if needed
                                        # Example: if child_el.get('id') == 'gap-fill-answers': ...
                                        div_text_content = child_el.get_text(separator='\n', strip=True)
                                        if div_text_content: document.add_paragraph(div_text_content)
                            # Add more specific handlers for other common HTML structures here
                        document.add_paragraph() # Space after generic item
                    else:
                        logging.warning(f"Could not find content container for generic item {item.id}")
                        document.add_paragraph(f"[Content for item {item.id} (type: {item.item_type}) could not be extracted.]", style='Comment')
                except Exception as generic_err:
                    logging.error(f"Error processing generic text-based item {item.id}: {generic_err}", exc_info=True)
                    document.add_paragraph(f"[Error processing item {item.id} (type: {item.item_type})]", style='Comment')

            logging.debug(f"End processing item index {item_index}. Identified as: {processed_as_type}")
            if item_index < len(items) - 1: # If it's not the last item
                # Add a page break only before large visual items if the current item wasn't also one
                # or if the next item is a major section.
                # For now, simpler: page break before WordSearch or ImageLabel if they aren't first.
                next_item_is_visual = items[item_index + 1].item_type in ['wordSearch', 'imageLabel']
                current_item_is_visual = item.item_type in ['wordSearch', 'imageLabel']

                if next_item_is_visual and item_index > -1 : # Always page break before a visual item unless it's the very first
                    logging.debug(f"Adding PAGE BREAK before next visual item: {items[item_index + 1].item_type}")
                    document.add_page_break()
                elif not current_item_is_visual and not next_item_is_visual: 
                    # If both current and next are text-based, add a smaller separator
                    logging.debug(f"Adding paragraph separator after text item: {item.item_type}")
                    sep_para = document.add_paragraph()
                    # You can add a run with '***' or a few empty runs, or just rely on paragraph spacing.
                    # sep_para.add_run("-----").font.size = Pt(8) 
                    # sep_para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    document.add_paragraph() # Just an empty paragraph for spacing
                else:
                    # If current is visual and next is text, or vice-versa and not covered above,
                    # default to a paragraph break.
                    logging.debug(f"Adding default paragraph separator between {item.item_type} and {items[item_index + 1].item_type}")
                    document.add_paragraph()

        logging.info("Finished item processing loop.")
        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0)
        safe_title = re.sub(r'[\\/*?:"<>|]', "", worksheet.title)
        safe_title = re.sub(r'\s+', '_', safe_title).strip('_')
        filename = f"{safe_title.lower() or 'worksheet'}.docx"
        logging.info(f"Sending DOCX file: {filename}")
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        logging.error(f"CRITICAL Error exporting DOCX for worksheet {worksheet_id}", exc_info=True)
        return f"Error exporting worksheet: {str(e)}<br><pre>{traceback.format_exc()}</pre>", 500

# --- Helper function add_runs_from_html_element ---
def add_runs_from_html_element(paragraph, element):
    """Adds formatted runs to a python-docx paragraph based on basic HTML tags."""
    for content in element.contents:
        if isinstance(content, str):
            text_content_original = content 
            text_content_for_display = content.replace('\xa0', ' ') 

            stripped_text = text_content_original.strip()
            is_just_underscores = False
            if len(stripped_text) > 1 and all(char == '_' for char in stripped_text): 
                is_just_underscores = True
            
            if is_just_underscores:
                paragraph.add_run(" " + "_" * 30 + " ") # Make the gap 30 underscores long
            else:
                text_to_add_final = text_content_for_display.strip()
                if text_to_add_final:
                    paragraph.add_run(text_to_add_final)
        elif content.name in ['strong', 'b']:
            run = paragraph.add_run(content.get_text(strip=True).replace('\xa0', ' '))
            run.bold = True
        elif content.name in ['em', 'i']:
            run = paragraph.add_run(content.get_text(strip=True).replace('\xa0', ' '))
            run.italic = True
        elif content.name == 'span' and 'gap-placeholder' in content.get('class', []): 
            # This will now only be hit if you explicitly use <span class="gap-placeholder"> for some gaps
            paragraph.add_run(" [____________________] ") # Keep it long here too
        elif content.name == 'br':
            paragraph.add_run().add_break() 
        elif hasattr(content, 'name') and content.name: 
            nested_text = content.get_text(strip=True).replace('\xa0', ' ')
            if nested_text:
                paragraph.add_run(nested_text)

@app.route('/')
def serve_index():
    """Serves the index.html file."""
    logging.info("Serving index.html")
    return send_from_directory('.', 'index.html')

# --- Run the App ---
if __name__ == '__main__':
    # db.create_all() is now handled when the app initializes globally
    app.run(debug=True, port=5001)
